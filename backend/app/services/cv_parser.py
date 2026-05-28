"""
CV text extraction from PDF/DOCX files.

Parsing logic lives in mock_ai.py (default) and ai_client.py (optional live APIs).
"""

import logging
from pathlib import Path

from app.services.ai_client import AIClient
from app.services.mock_ai import enrich_parsed_cv, mock_extract_cv
from app.utils.text_normalize import normalize_extracted_text

logger = logging.getLogger(__name__)


def extract_text_from_file(file_path: Path) -> str:
    """Read plain text from an uploaded PDF or DOCX file with spacing normalization."""
    suffix = file_path.suffix.lower()
    raw = ""

    if suffix == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(str(file_path))
        pages = [page.extract_text() or "" for page in reader.pages]
        raw = "\n".join(pages).strip()
        if not raw:
            raise ValueError(
                "Could not extract text from PDF. "
                "The file may be scanned/image-only — use a text-based PDF or DOCX."
            )

    elif suffix == ".docx":
        from docx import Document

        doc = Document(str(file_path))
        parts: list[str] = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        raw = "\n".join(parts).strip()
        if not raw:
            raise ValueError("Could not extract text from DOCX — the document appears empty.")

    elif suffix == ".doc":
        raise ValueError(
            "Legacy .doc format is not supported. Please save your CV as .docx or export to PDF."
        )
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    normalized = normalize_extracted_text(raw)
    logger.info(
        "Extracted text from %s: raw_len=%s normalized_len=%s",
        file_path.name,
        len(raw),
        len(normalized),
    )
    return normalized


async def parse_cv_with_ai(text: str, ai: AIClient) -> dict:
    """Try live AI extraction first; always merge heuristic skills before return."""
    result = await ai.extract_cv_profile(text)
    if result:
        result.setdefault("extraction_method", "live_ai+heuristic")
        return enrich_parsed_cv(result, text)
    return mock_extract_cv(text)
